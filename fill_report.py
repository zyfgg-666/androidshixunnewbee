import os, glob
os.chdir(r"D:\soft\android202422\workspace\NewBee2")

template = None
for f in glob.glob("*实训报告模板*"):
    template = f; break

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

doc = Document(template)

# ====== Fill paragraphs ======
# P24: 项目成员分工
doc.paragraphs[24].text = ""  # clear, we'll add content after

# P25: System design
doc.paragraphs[26].text = """2.2 系统设计与实现

本App采用MVC架构，主页面使用"单Activity + ViewPager2 + 4个Fragment + BottomNavigationView"架构。
ViewPager2管理4个Fragment（首页/分类/购物车/我的），底部导航栏与之双向联动。

技术栈：Java 11 + HttpURLConnection(线程池异步) + Gson 2.10.1 + Glide 4.12.0

网络层封装HttpUtil类，提供GET/POST/PUT/DELETE五种HTTP方法，自动注入Token认证头，
通过Handler+Looper切回主线程回调。Token存储在SharedPreferences中。

共对接19个RESTful API接口，分为用户、商品、购物车、订单、地址5个模块。"""

# P28: 概要设计
doc.paragraphs[28].text = """2.3 概要设计

功能模块划分：
(1) 启动页：品牌Logo展示，2秒自动跳转
(2) 登录注册：账号密码登录，MD5加密，Token认证
(3) 首页：轮播图(3秒自动轮播)、分类网格(4列)、三大商品板块(2列)
(4) 分类：左右联动、首页分类点击自动选中
(5) 搜索：关键词搜索、排序(默认/新品/价格)、上拉分页
(6) 详情：大图+标签+原价、加购+立即购买
(7) 购物车：全选联动、数量修改、删除、结算过滤
(8) 订单：创建订单、支付(微信/支付宝/稍后支付)、6Tab列表、确认收货
(9) 地址：增删改查、省市区三级联动、默认地址
(10) 我的：用户信息、订单红点、昵称签名编辑、账号管理

页面导航：
SplashActivity → MainActivity(ViewPager2)
  ├── HomeFragment → SearchActivity / DetailActivity
  ├── CategoryFragment → DetailActivity
  ├── CartFragment → CreateOrderActivity → OrderListActivity
  └── MyFragment → OrderListActivity / AddressListActivity / AboutActivity"""

# P30: 详细设计
doc.paragraphs[30].text = """2.4 详细设计及实现（按分工详述）"""

# Now fill the detailed content after P30
# We need to add detailed paragraphs. Let me insert text after the section headers.

# ---- 首页 ----
p = doc.add_paragraph()
p.style = doc.styles['Normal']
run = p.add_run("""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.1 首页（张玉峰 542407230733）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：轮播图(ViewPager2自动轮播，onPause暂停)、搜索栏(青绿色背景)、分类网格
(MyGridView 4列)、三大商品板块(新品/热销/推荐)。核心技术：Handler+Runnable
定时器实现3秒轮播；自定义MyGridView重写onMeasure解决ScrollView嵌套GridView
高度问题；/index-infos接口一次返回4种数据；分类点击通过CategoryFragment.
setPendingCategory()+switchToTab()实现跳转并选中。
文件：HomeFragment.java(298行)、fragment_home.xml(415行)、BannerAdapter.java(50行)、
CategoryGridAdapter.java(92行)、GoodsAdapter.java(76行)、MyGridView.java(27行)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.2 购物车（张玉峰 542407230733）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：RecyclerView列表、全选联动(防递归)、数量±同步服务器、删除、空状态(图标+
"快去选购")。核心技术：全选CheckBox与单选CheckBox双向联动，checkAllSelect()
采用"取消监听→setChecked→恢复监听"防递归；结算时通过Intent传递选中cartItemId
数组给CreateOrderActivity过滤；Material3主题下Button改用TextView+drawable背景。
文件：CartFragment.java(269行)、CartAdapter.java(116行)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.3 商品搜索（张玉峰 542407230733）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：关键词搜索、排序切换(默认/新品/价格，青绿色高亮)、上拉加载更多(RecyclerView
OnScrollListener检测到底部→page++追加数据)、2列网格展示。
文件：SearchActivity.java(149行)、SearchGoodsAdapter.java(73行)""")

# ---- 商品详情 ----
p = doc.add_paragraph()
run = p.add_run("""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.4 商品详情（张玉峰 542407230733 / 吴家乐 542407230725）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：商品大图(300dp)、名称+简介+标签(橙色底)+现价(红)+原价(灰色删除线)、
购物车红点角标(FrameLayout+TextView 14dp)、加购(POST /shop-cart)、立即购买
(传递singleGoodsId→CreateOrderActivity，仅含该商品不混入购物车其他商品)。
文件：DetailActivity.java(226行)、activity_detail.xml(172行)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.5 分类（吴家乐 542407230725）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：青绿标题栏、左侧一级分类RecyclerView(选中白底青字高亮)、右侧2列GridView
商品网格。核心技术：分类名→搜索关键词switch-case映射(家电→手机/女装→T恤等)，
因后端goodsCategoryId参数不生效的变通方案；静态pendingCategoryId实现首页→分类
自动选中。
文件：CategoryFragment.java(276行)、CategoryLeftAdapter.java(82行)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.6 生成订单 + 支付（吴家乐 542407230725）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：加载购物车→选择/自动加载默认地址→POST /saveOrder→显示订单号→弹支付对话框
(微信绿色/支付宝蓝色/稍后支付灰色)。核心技术：startActivityForResult接收地址选择；
GET /paySuccess?orderNo=xxx&payType=1/2模拟支付；稍后支付不调接口直接跳转待支付。
文件：CreateOrderActivity.java(281行)、dialog_pay.xml""")

# ---- 我的 ----
p = doc.add_paragraph()
run = p.add_run("""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.7 我的（肖奥涵 542407230726）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：青绿渐变用户卡片(头像+昵称+签名)、6个订单入口图标(待支付/待发货/待收货
红点角标)、昵称/签名点击编辑(AlertDialog+EditText→PUT /user/info)、账号管理
(查看登录账号+修改密码PUT /user/password)、退出登录。
文件：MyFragment.java(319行)、fragment_my.xml(约400行)、AccountManageActivity.java

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.8 地址管理（肖奥涵 542407230726）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：地址列表(普通/选择双模式)、新增/编辑(收货人+手机+省市区+详情)、省市区三级
联动选择器(三栏ListView+31省数据+选中红字高亮)、默认地址设置(PUT /address完整字段)。
核心技术：RegionData硬编码31省数据；三栏联动：选省→更新城市列表→选市→更新区县。
文件：AddressListActivity.java(180行)、AddressEditActivity.java(311行)、
AddressAdapter.java(107行)、RegionData.java(183行)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2.4.9 我的订单（肖奥涵 542407230726）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

功能：6个Tab(全部/待支付/待确认/待发货/待收货/已完成)、订单卡片(订单号+商品
嵌套RecyclerView+状态+合计)、待支付→"立即支付"按钮、待收货→"确认收货"按钮。
核心技术：TabLayout+RecyclerView筛选；订单状态流转0→(支付)→2→(发货)→3→(收货)→4；
嵌套RecyclerView设置clickable=false避免点击穿透。
文件：OrderListActivity.java(133行)、OrderAdapter.java(257行)、
OrderGoodsAdapter.java(62行)""")

# ---- P33: 实训体会 ----
doc.paragraphs[33].text = ""  # clear header, content will be added below

p = doc.add_paragraph()
run = p.add_run("""
【张玉峰】通过本次实训，我深入了解了Android原生开发的核心技术栈。在首页开发中，
自定义MyGridView解决ScrollView嵌套GridView高度问题的经历让我理解了Android的
MeasureSpec机制。购物车全选防递归问题的解决（取消监听→设值→恢复监听）是经典
的Android CheckBox处理方案。搜索模块的分页加载让我掌握了RecyclerView滚动监听。
最大的收获是学会了前后端协作的完整流程和Git版本管理。

【吴家乐】在分类模块中，因后端goodsCategoryId参数不生效，我采用了分类名→搜索
关键词的switch-case映射方案，学会了在后端限制下寻找替代方案。支付接口调试是最大
收获：最初返回500错误，通过查阅Swagger文档和curl测试，发现缺少payType参数。
这个过程让我掌握了API调试的标准流程：查文档→curl验证→修改代码。

【肖奥涵】地址管理的省市区三级联动选择器是最具挑战性的任务。我硬编码了31个省份
数据到RegionData类，使用三栏ListView+自定义Adapter实现选中高亮。订单列表的
TabLayout筛选和订单状态流转让我理解了电商业务流程。昵称签名编辑通过AlertDialog+
PUT接口实现在线修改。实训让我全面掌握了Fragment开发、RecyclerView适配器模式和
SharedPreferences数据存储。""")

# ====== Fill Table 0: 实训日志 ======
logs = [
    "实训动员，环境搭建：安装Android Studio，配置JDK/Gradle，创建项目，配置OkHttp/Gson/Glide等依赖",
    "搭建项目框架：ViewPager2+BottomNavigationView主页结构，编写4个Fragment骨架，封装HttpUtil网络层(HttpURLConnection+线程池)",
    "完成登录注册模块：LoginActivity，MD5加密，Token存储到SharedPreferences，登录接口联调测试",
    "首页开发：轮播图(ViewPager2自动轮播，Handler+Runnable定时器)，分类网格(4列MyGridView)，三大商品板块(新品/热销/推荐)，自定义MyGridView解决ScrollView嵌套GridView高度问题",
    "分类页开发：左右联动(RecyclerView+GridView)；搜索页开发：关键词搜索、排序切换(默认/新品/价格)、上拉分页；商品详情页：大图+标签+原价、加购+立即购买",
    "购物车开发：RecyclerView列表、全选联动(防递归方案)、数量修改同步服务器(POST /shop-cart/update)、删除、结算过滤(传选中cartItemId数组)、空购物车状态(图标+快去选购按钮)、底部导航角标(BadgeDrawable)",
    "订单系统开发：创建订单(POST /saveOrder)、支付对话框(微信/支付宝/稍后支付)、支付接口GET /paySuccess联调(修复缺少payType参数导致500错误)、订单列表6Tab、确认收货(PUT /order/{no}/finish)",
    "地址管理：增删改查、省市区三级联动选择器(三栏ListView+31省RegionData数据+选中红字高亮)、默认地址设置(PUT发送完整字段)；我的页：用户信息卡片、订单红点角标、昵称签名编辑弹窗(AlertDialog+PUT /user/info)",
    "UI美化：30+矢量图标设计(底部导航/订单/分类/功能入口)、颜色体系统一(#00C9A7青绿)、Material3主题下Button→TextView样式修复、启动页(蜜蜂手推购物车)+App图标(蜜蜂拎购物袋)设计、关于我们页面、网络层从OkHttp改为HttpURLConnection(实训要求)",
    "Bug修复+测试+答辩准备：购物车结算选中过滤、立即购买单个商品ID传递、默认地址自动加载到确认订单页、账号管理功能、商品详情完善(标签+原价)；全面功能测试；编写答辩材料(项目答辩材料.txt/代码汇总/网络通信详解/首页实现详解/购物车实现详解等6份文档)"
]

table = doc.tables[0]
for i in range(min(len(logs), len(table.rows)-1)):
    row = table.rows[i+1]
    row.cells[1].text = logs[i]

# ====== Fill Table 1: 实训成绩 ======
members = [
    ("542407230733", "张玉峰", "首页、购物车、商品搜索、商品详情"),
    ("542407230725", "吴家乐", "分类、生成订单、商品详情、支付"),
    ("542407230726", "肖奥涵", "我的、地址管理、我的订单"),
]
table2 = doc.tables[1]
for i, (sid, name, work) in enumerate(members):
    row = table2.rows[i+1]
    row.cells[0].text = sid
    row.cells[1].text = name
    row.cells[2].text = work

# ====== Save ======
outpath = "移动应用程序设计实训报告-已填写.docx"
doc.save(outpath)
print(f"Saved to: {outpath}")
